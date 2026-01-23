import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_guide():
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Guía de Migración y Despliegue en VPS - ERP EDUCATIVA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Objetivo: Desplegar el sistema completo (Frontend React + Backend Django + Base de Datos PostgreSQL) en un servidor VPS Linux (Ubuntu 22.04/24.04) desde cero.')
    doc.add_paragraph('Este documento está diseñado para ser seguido paso a paso. Puede copiar y pegar los comandos directamente en la terminal de su servidor.')

    # Function to add code block
    def add_code_block(code_text):
        p = doc.add_paragraph()
        p.style = 'No Spacing'
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 100) # Dark Blue
        p.paragraph_format.left_indent = Inches(0.5)

    # 1. Prerrequisitos
    doc.add_heading('1. Prerrequisitos y Acceso al Servidor', level=1)
    doc.add_paragraph('Antes de comenzar, asegúrese de tener:')
    p = doc.add_paragraph()
    p.add_run('1. Acceso SSH al VPS:').bold = True
    p.add_run(' IP del servidor, usuario (usualmente root) y contraseña.')
    p = doc.add_paragraph()
    p.add_run('2. Dominio (Opcional pero recomendado):').bold = True
    p.add_run(' Un dominio (ej: mi-erp.com) apuntando a la IP del VPS.')
    p = doc.add_paragraph()
    p.add_run('3. Cliente SSH:').bold = True
    p.add_run(' Putty (Windows) o Terminal (Mac/Linux).')

    doc.add_heading('Paso 1.1: Conexión al Servidor', level=2)
    doc.add_paragraph('Abra su terminal o Putty y conéctese:')
    add_code_block("ssh root@TU_IP_DEL_SERVIDOR")
    doc.add_paragraph('(Si es la primera vez, acepte la huella digital escribiendo yes).').italic = True

    # 2. Preparación
    doc.add_heading('2. Preparación del Sistema Operativo', level=1)
    doc.add_paragraph('Actualizaremos el sistema e instalaremos las herramientas base.')
    
    doc.add_heading('Paso 2.1: Actualizar Repositorios', level=2)
    add_code_block("sudo apt update && sudo apt upgrade -y")

    doc.add_heading('Paso 2.2: Instalar Dependencias del Sistema', level=2)
    doc.add_paragraph('Instalaremos Python, PostgreSQL, Redis, Nginx, Git y herramientas de compilación.')
    add_code_block("""# Herramientas básicas y compiladores
sudo apt install -y git curl wget build-essential libssl-dev libffi-dev

# Dependencias para PostgreSQL y Python
sudo apt install -y python3-pip python3-venv python3-dev libpq-dev postgresql postgresql-contrib

# Servidor Web y Redis (para WebSockets/Cache)
sudo apt install -y nginx redis-server

# Firewall (UFW)
sudo ufw allow OpenSSH
sudo ufw allow 'Nginx Full'
sudo ufw enable""")
    
    doc.add_heading('Paso 2.3: Instalar Node.js', level=2)
    add_code_block("""curl -fsSL https://deb.nodesource.com/setup_20.x | sudo -E bash -
sudo apt install -y nodejs""")

    # 3. Base de Datos
    doc.add_heading('3. Configuración de la Base de Datos (PostgreSQL)', level=1)
    doc.add_heading('Paso 3.1: Crear Base de Datos y Usuario', level=2)
    add_code_block("sudo -u postgres psql")
    doc.add_paragraph('Dentro de la consola de PostgreSQL:')
    add_code_block("""CREATE DATABASE erp_educativa;
CREATE USER erp_admin WITH PASSWORD 'SecretoSeguro123';
ALTER ROLE erp_admin SET client_encoding TO 'utf8';
ALTER ROLE erp_admin SET default_transaction_isolation TO 'read committed';
ALTER ROLE erp_admin SET timezone TO 'UTC';
GRANT ALL PRIVILEGES ON DATABASE erp_educativa TO erp_admin;
GRANT ALL ON SCHEMA public TO erp_admin;
\\q""")

    # 4. Backend
    doc.add_heading('4. Despliegue del Backend (Django)', level=1)
    doc.add_heading('Paso 4.1: Estructura de Directorios', level=2)
    add_code_block("""sudo mkdir -p /var/www/erp_educativa
sudo chown -R $USER:www-data /var/www/erp_educativa
sudo chmod -R 775 /var/www/erp_educativa
cd /var/www/erp_educativa""")

    doc.add_heading('Paso 4.2: Obtener el Código', level=2)
    doc.add_paragraph('Opción A (Git):')
    add_code_block("git clone <URL_DE_SU_REPOSITORIO> .")
    
    doc.add_heading('Paso 4.3: Configurar Entorno Virtual', level=2)
    add_code_block("""cd /var/www/erp_educativa/backend
python3 -m venv venv
source venv/bin/activate
pip install --upgrade pip
pip install -r requirements.txt
pip install gunicorn daphne psycopg2-binary""")

    doc.add_heading('Paso 4.4: Configurar Variables de Entorno (.env)', level=2)
    doc.add_paragraph('Cree el archivo .env en backend/ con nano:')
    add_code_block("""DEBUG=False
SECRET_KEY=cambiar_por_una_clave_larga_y_aleatoria
ALLOWED_HOSTS=tudominio.com,www.tudominio.com,TU_IP_PUBLICA
DB_ENGINE=django.db.backends.postgresql
DB_NAME=erp_educativa
DB_USER=erp_admin
DB_PASSWORD=SecretoSeguro123
DB_HOST=localhost
DB_PORT=5432
REDIS_URL=redis://127.0.0.1:6379/1""")

    doc.add_heading('Paso 4.5: Migraciones y Estáticos', level=2)
    add_code_block("""python manage.py collectstatic --noinput
python manage.py migrate
python manage.py createsuperuser""")

    doc.add_heading('Paso 4.6: Configurar Servicios Systemd', level=2)
    
    doc.add_paragraph('1. Gunicorn (HTTP):')
    doc.add_paragraph('/etc/systemd/system/erp_gunicorn.service').bold = True
    add_code_block("""[Unit]
Description=Gunicorn daemon for ERP Educativa
After=network.target

[Service]
User=root
Group=www-data
WorkingDirectory=/var/www/erp_educativa/backend
EnvironmentFile=/var/www/erp_educativa/backend/.env
ExecStart=/var/www/erp_educativa/backend/venv/bin/gunicorn --access-logfile - --workers 3 --bind unix:/run/erp_gunicorn.sock config.wsgi:application

[Install]
WantedBy=multi-user.target""")

    doc.add_paragraph('2. Daphne (WebSockets):')
    doc.add_paragraph('/etc/systemd/system/erp_daphne.service').bold = True
    add_code_block("""[Unit]
Description=Daphne daemon for ERP WebSockets
After=network.target

[Service]
User=root
Group=www-data
WorkingDirectory=/var/www/erp_educativa/backend
EnvironmentFile=/var/www/erp_educativa/backend/.env
ExecStart=/var/www/erp_educativa/backend/venv/bin/daphne -b 0.0.0.0 -p 8001 config.asgi:application

[Install]
WantedBy=multi-user.target""")

    doc.add_paragraph('3. Celery (Worker):')
    doc.add_paragraph('/etc/systemd/system/erp_celery.service').bold = True
    add_code_block("""[Unit]
Description=Celery Worker for ERP Educativa
After=network.target redis-server.service

[Service]
User=root
Group=www-data
WorkingDirectory=/var/www/erp_educativa/backend
EnvironmentFile=/var/www/erp_educativa/backend/.env
ExecStart=/var/www/erp_educativa/backend/venv/bin/celery -A config worker -l info

[Install]
WantedBy=multi-user.target""")

    doc.add_paragraph('Iniciar servicios:')
    add_code_block("""sudo systemctl daemon-reload
sudo systemctl start erp_gunicorn erp_daphne erp_celery
sudo systemctl enable erp_gunicorn erp_daphne erp_celery""")

    doc.add_heading('Paso 4.7: Permisos Finales', level=2)
    add_code_block("""sudo chown -R :www-data /var/www/erp_educativa/backend/media
sudo chown -R :www-data /var/www/erp_educativa/backend/static
sudo chmod -R 775 /var/www/erp_educativa/backend/media
sudo chmod -R 775 /var/www/erp_educativa/backend/static""")

    # 5. Frontend
    doc.add_heading('5. Despliegue del Frontend (React)', level=1)
    add_code_block("""cd /var/www/erp_educativa/frontend
npm install
# Crear .env.production con VITE_API_URL=https://tudominio.com/api
npm run build""")

    # 6. Nginx
    doc.add_heading('6. Configuración de Nginx', level=1)
    doc.add_paragraph('/etc/nginx/sites-available/erp_educativa')
    add_code_block("""server {
    listen 80;
    server_name tudominio.com www.tudominio.com;

    location / {
        root /var/www/erp_educativa/frontend/dist;
        try_files $uri $uri/ /index.html;
    }

    location /api/ {
        include proxy_params;
        proxy_pass http://unix:/run/erp_gunicorn.sock;
    }

    location /admin/ {
        include proxy_params;
        proxy_pass http://unix:/run/erp_gunicorn.sock;
    }

    location /static/ {
        alias /var/www/erp_educativa/backend/static/;
    }

    location /media/ {
        alias /var/www/erp_educativa/backend/media/;
    }

    location /ws/ {
        proxy_pass http://127.0.0.1:8001;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection "upgrade";
    }
}""")
    doc.add_paragraph('Activar sitio:')
    add_code_block("""sudo ln -s /etc/nginx/sites-available/erp_educativa /etc/nginx/sites-enabled/
sudo nginx -t
sudo systemctl restart nginx""")

    # 7. SSL
    doc.add_heading('7. Seguridad SSL (Certbot)', level=1)
    add_code_block("""sudo apt install certbot python3-certbot-nginx
sudo certbot --nginx -d tudominio.com""")

    # 8. Troubleshooting
    doc.add_heading('8. Solución de Problemas Comunes', level=1)
    doc.add_paragraph('Error 502 Bad Gateway:', style='List Bullet')
    doc.add_paragraph('Verificar Gunicorn: sudo systemctl status erp_gunicorn', style='List Bullet 2')
    doc.add_paragraph('Error de Base de Datos:', style='List Bullet')
    doc.add_paragraph('Verificar .env y permisos de usuario PostgreSQL.', style='List Bullet 2')
    doc.add_paragraph('Archivos Estáticos 403:', style='List Bullet')
    doc.add_paragraph('Revisar permisos chown/chmod explicados anteriormente.', style='List Bullet 2')

    # Save
    output_path = os.path.join(os.getcwd(), 'Guia_Migracion_VPS.docx')
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

if __name__ == "__main__":
    create_word_guide()
