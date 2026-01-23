from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    document = Document()

    # Title
    title = document.add_heading('Manual Operacional del Sistema ERP EDUCATIVA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Introducción
    document.add_heading('1. Introducción', level=1)
    document.add_paragraph(
        'ERP EDUCATIVA es un sistema integral de gestión para instituciones educativas que cubre áreas académicas, '
        'administrativas, financieras y de comunicación. Este manual detalla los procedimientos operativos estándar '
        'para administradores, docentes y personal administrativo.'
    )

    # 2. Acceso y Puesta en Marcha
    document.add_heading('2. Acceso y Puesta en Marcha', level=1)

    document.add_heading('2.1 Requisitos del Sistema', level=2)
    p = document.add_paragraph()
    p.add_run('Servidor: ').bold = True
    p.add_run('Docker instalado y ejecutándose.')
    p = document.add_paragraph()
    p.add_run('Cliente: ').bold = True
    p.add_run('Navegador web moderno (Chrome, Firefox, Edge).')

    document.add_heading('2.2 Inicio del Sistema', level=2)
    document.add_paragraph(
        'El sistema se ejecuta mediante contenedores Docker y un servidor de desarrollo local (mientras no esté en producción).'
    )

    document.add_paragraph('1. Iniciar Base de Datos (PostgreSQL):', style='List Number')
    document.add_paragraph('docker-compose up -d db', style='Quote')

    document.add_paragraph('2. Iniciar Backend (Django):', style='List Number')
    document.add_paragraph('python backend/manage.py runserver 0.0.0.0:8000', style='Quote')
    document.add_paragraph('El servidor estará disponible en http://localhost:8000', style='Intense Quote')

    document.add_paragraph('3. Iniciar Frontend (React):', style='List Number')
    document.add_paragraph('cd frontend\nnpm run dev', style='Quote')
    document.add_paragraph('La aplicación estará disponible en http://localhost:5173 (o el puerto que indique la consola)', style='Intense Quote')

    # 3. Gestión de Usuarios y Roles
    document.add_heading('3. Gestión de Usuarios y Roles', level=1)

    document.add_heading('3.1 Roles Disponibles', level=2)
    roles = [
        ('Admin:', ' Acceso total al sistema.'),
        ('Rector:', ' Gestión académica y reportes.'),
        ('Profesor:', ' Gestión de cursos asignados, calificaciones y asistencia.'),
        ('Estudiante/Padre:', ' Consulta de notas, asistencia y comunicación.')
    ]
    for role, desc in roles:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(role).bold = True
        p.add_run(desc)

    document.add_heading('3.2 Creación de Usuarios', level=2)
    steps_users = [
        'Navegue a Usuarios en el menú principal.',
        'Haga clic en "Crear Usuario".',
        'Complete los campos obligatorios: Nombre, Usuario, Contraseña, Rol e Institución.',
        'Nota: El usuario admin por defecto se crea con la institución "INSTITUCION PRUEBA".'
    ]
    for step in steps_users:
        document.add_paragraph(step, style='List Number')

    # 4. Módulo Académico
    document.add_heading('4. Módulo Académico', level=1)

    document.add_heading('4.1 Configuración Inicial (Inicio de Año)', level=2)
    
    document.add_paragraph('1. Crear Año Lectivo:', style='List Number')
    document.add_paragraph('Vaya a Años Lectivos.', style='List Bullet 2')
    document.add_paragraph('Cree un nuevo año (ej: "2025-2026").', style='List Bullet 2')
    document.add_paragraph('El sistema generará automáticamente los periodos (Trimestres).', style='List Bullet 2')

    document.add_paragraph('2. Apertura de Periodos:', style='List Number')
    document.add_paragraph('Active el año lectivo actual.', style='List Bullet 2')
    document.add_paragraph('Asegúrese de que el periodo actual (ej: Trimestre 1) tenga el candado ABIERTO.', style='List Bullet 2')
    p = document.add_paragraph(style='List Bullet 2')
    p.add_run('Importante: ').bold = True
    p.add_run('Si un periodo está cerrado, los profesores NO podrán ingresar o modificar notas.')

    document.add_heading('4.2 Gestión de Cursos y Materias', level=2)
    course_steps = [
        ('Cursos:', ' Defina los niveles (ej: "1ro Bachillerato A").'),
        ('Materias:', ' Cree asignaturas (ej: "Matemáticas", "Historia").'),
        ('Asignación:', ' Asigne materias a cursos y defina el profesor encargado.')
    ]
    for title, desc in course_steps:
        p = document.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(desc)

    document.add_heading('4.3 Matrículas', level=2)
    document.add_paragraph('1. Inscriba estudiantes en los cursos correspondientes desde la sección Matrículas.', style='List Number')

    document.add_heading('4.4 Ingreso de Calificaciones (Profesores)', level=2)
    grade_steps = [
        'El profesor ingresa a Mis Cursos.',
        'Selecciona la materia y luego Calificaciones.',
        'Ingresa las notas en las categorías correspondientes (Aportes, Examen, etc.).',
        'Nota: El sistema valida que la nota esté entre 0 y 10.'
    ]
    for step in grade_steps:
        document.add_paragraph(step, style='List Number')

    # 5. Mantenimiento del Sistema
    document.add_heading('5. Mantenimiento del Sistema (Danger Zone)', level=1)

    document.add_heading('5.1 Respaldo (Backup)', level=2)
    document.add_paragraph('1. Puede descargar un volcado (dump) de la base de datos desde la opción de mantenimiento (si está habilitada).', style='List Number')

    document.add_heading('5.2 Reseteo del Sistema (Reset System)', level=2)
    p = document.add_paragraph()
    p.add_run('¡ADVERTENCIA! ').bold = True
    p.add_run('Esta acción es destructiva.').italic = True

    reset_steps = [
        'Navegue a Mantenimiento > System Reset.',
        'Escriba "RESETEAR" para confirmar.',
        'El sistema borrará TODOS los datos académicos, contables y de usuarios.',
        'Recreará automáticamente su usuario admin actual para que no pierda acceso.',
        'Reiniciará la institución a "INSTITUCION PRUEBA".'
    ]
    for step in reset_steps:
        document.add_paragraph(step, style='List Number')

    # 6. Solución de Problemas
    document.add_heading('6. Solución de Problemas Comunes', level=1)

    document.add_heading('6.1 "No puedo ingresar notas"', level=2)
    p = document.add_paragraph()
    p.add_run('Causa Probable: ').bold = True
    p.add_run('El periodo académico o el año lectivo están cerrados.')
    document.add_paragraph('Solución: Un administrador debe ir a Años Lectivos y abrir el candado del periodo correspondiente.')

    document.add_heading('6.2 "Error de Conexión"', level=2)
    p = document.add_paragraph()
    p.add_run('Causa Probable: ').bold = True
    p.add_run('Docker no está corriendo o el puerto 8000 está ocupado.')
    document.add_paragraph('Solución:')
    document.add_paragraph('1. Reinicie Docker.', style='List Number')
    document.add_paragraph('2. Reinicie el servidor Django.', style='List Number')
    document.add_paragraph('3. Verifique que no haya otras instancias de Python corriendo.', style='List Number')

    document.add_heading('Soporte Técnico', level=1)
    document.add_paragraph('Para errores críticos, contacte al administrador del sistema o consulte los logs en el backend.')

    document.save('Manual_Operacional_ERP.docx')
    print("Manual creado exitosamente: Manual_Operacional_ERP.docx")

if __name__ == "__main__":
    create_manual()
